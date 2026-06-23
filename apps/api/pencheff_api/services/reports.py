"""Generate downloadable reports in DOCX / CSV / JSON / PDF."""
from __future__ import annotations

import asyncio
import csv
import io
import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from ..config import get_settings
from ..db.models import Engagement, Finding, Report, Scan, Target, WorkspaceBranding

log = logging.getLogger(__name__)

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
SEVERITY_HEX = {
    "critical": "C00000", "high": "E06666", "medium": "E69138",
    "low": "6FA8DC", "info": "B7B7B7",
}


def _storage_dir() -> Path:
    p = Path(get_settings().report_storage_dir)
    p.mkdir(parents=True, exist_ok=True)
    return p


def _timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


def _sort_findings(findings: list[Finding]) -> list[Finding]:
    return sorted(findings, key=lambda f: (SEVERITY_ORDER.get((f.severity or "info").lower(), 99), f.created_at))


# ── JSON ──────────────────────────────────────────────────────────────

def write_json(scan: Scan, target: Target, findings: list[Finding], path: Path) -> int:
    data = {
        "report_metadata": {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "target": target.base_url,
            "target_name": target.name,
            "scan_id": scan.id,
            "profile": scan.profile,
            "grade": scan.grade,
            "score": scan.score,
        },
        "summary": scan.summary or {},
        "findings": [
            {
                "id": f.id, "title": f.title, "severity": f.severity,
                "category": f.category, "owasp_category": f.owasp_category,
                "cwe_id": f.cwe_id, "cvss_score": f.cvss_score, "cvss_vector": f.cvss_vector,
                "endpoint": f.endpoint, "parameter": f.parameter,
                "description": f.description, "remediation": f.remediation,
                "evidence": f.evidence, "references": f.references_,
                "verification_status": f.verification_status,
                "suppressed": f.suppressed, "suppress_reason": f.suppress_reason,
                "last_rechecked_at": f.last_rechecked_at.isoformat() if f.last_rechecked_at else None,
            }
            for f in _sort_findings(findings)
        ],
    }
    payload = json.dumps(data, indent=2, default=str)
    path.write_text(payload, encoding="utf-8")
    return len(payload.encode())


# ── CSV ───────────────────────────────────────────────────────────────

def write_csv(scan: Scan, target: Target, findings: list[Finding], path: Path) -> int:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["id", "title", "severity", "category", "owasp", "cwe", "cvss_score",
                "endpoint", "parameter", "verification_status", "suppressed",
                "description", "remediation"])
    for f in _sort_findings(findings):
        w.writerow([
            f.id, f.title, f.severity, f.category, f.owasp_category or "",
            f.cwe_id or "", f.cvss_score or "", f.endpoint or "", f.parameter or "",
            f.verification_status, "yes" if f.suppressed else "no",
            (f.description or "").replace("\n", " ")[:2000],
            (f.remediation or "").replace("\n", " ")[:2000],
        ])
    path.write_text(buf.getvalue(), encoding="utf-8")
    return path.stat().st_size


# ── DOCX ──────────────────────────────────────────────────────────────

def write_docx(scan: Scan, target: Target, findings: list[Finding], path: Path) -> int:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    title = doc.add_heading("Pencheff Security Assessment Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"Target: ").bold = True
    meta.add_run(f"{target.name} — {target.base_url}\n")
    meta.add_run(f"Scan ID: ").bold = True
    meta.add_run(f"{scan.id}\n")
    meta.add_run(f"Generated: ").bold = True
    meta.add_run(f"{datetime.now(timezone.utc).isoformat()}\n")
    meta.add_run(f"Grade: ").bold = True
    grade_run = meta.add_run(f"{scan.grade or '-'}  (score: {scan.score or 0}/100)\n")
    grade_run.bold = True
    grade_run.font.size = Pt(16)

    doc.add_heading("Executive Summary", level=1)
    summary = scan.summary or {}
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Critical"
    hdr[1].text = "High"
    hdr[2].text = "Medium"
    hdr[3].text = "Low"
    hdr[4].text = "Info"
    row = table.add_row().cells
    row[0].text = str(summary.get("critical", 0))
    row[1].text = str(summary.get("high", 0))
    row[2].text = str(summary.get("medium", 0))
    row[3].text = str(summary.get("low", 0))
    row[4].text = str(summary.get("info", 0))

    doc.add_heading("Findings", level=1)
    for f in _sort_findings(findings):
        h = doc.add_heading(f"[{f.severity.upper()}] {f.title}", level=2)
        try:
            color = SEVERITY_HEX.get((f.severity or "info").lower(), "000000")
            for run in h.runs:
                run.font.color.rgb = RGBColor.from_string(color)
        except Exception:
            pass
        p = doc.add_paragraph()
        p.add_run("Category: ").bold = True
        p.add_run(f"{f.category}   ")
        if f.owasp_category:
            p.add_run("OWASP: ").bold = True
            p.add_run(f"{f.owasp_category}   ")
        if f.cvss_score:
            p.add_run("CVSS: ").bold = True
            p.add_run(f"{f.cvss_score} ({f.cvss_vector or ''})   ")
        if f.endpoint:
            p.add_run("Endpoint: ").bold = True
            p.add_run(f"{f.endpoint}   ")
        if f.parameter:
            p.add_run("Parameter: ").bold = True
            p.add_run(f"{f.parameter}")
        if f.description:
            doc.add_paragraph().add_run("Description: ").bold = True
            doc.add_paragraph(f.description)
        if f.remediation:
            doc.add_paragraph().add_run("Remediation: ").bold = True
            doc.add_paragraph(f.remediation)
        doc.add_paragraph("")

    doc.add_heading("Compliance Mapping", level=1)
    doc.add_paragraph(
        "Findings in this report map to controls from OWASP Top 10 2021, SOC 2 (CC6/CC7), "
        "PCI-DSS 4.0, NIST 800-53, ISO 27001:2022, and HIPAA Security Rule. "
        "This report is suitable for attachment to SOC 2 Type II and similar audit evidence packages."
    )
    notices = doc.add_paragraph()
    notices.add_run(
        "OWASP and OWASP Top 10 are trademarks of the OWASP Foundation. "
        "Pencheff is not affiliated with or endorsed by OWASP. Other framework "
        "names referenced above are the property of their respective owners and "
        "are used for identification only."
    ).italic = True

    doc.save(str(path))
    return path.stat().st_size


# ── PDF ───────────────────────────────────────────────────────────────

def write_pdf(scan: Scan, target: Target, findings: list[Finding], path: Path) -> int:
    from weasyprint import HTML

    rows = []
    for f in _sort_findings(findings):
        sev = (f.severity or "info").lower()
        color = "#" + SEVERITY_HEX.get(sev, "888888")
        rows.append(f"""
        <div class="finding sev-{sev}">
          <h3 style="color:{color}">[{sev.upper()}] {(f.title or '').replace('<', '&lt;')}</h3>
          <p><b>Category:</b> {f.category}
             {('<b>OWASP:</b> ' + f.owasp_category) if f.owasp_category else ''}
             {('<b>CVSS:</b> ' + str(f.cvss_score)) if f.cvss_score else ''}</p>
          <p><b>Endpoint:</b> {f.endpoint or '-'} &nbsp; <b>Parameter:</b> {f.parameter or '-'}</p>
          <p><b>Description:</b> {(f.description or '').replace('<', '&lt;')}</p>
          <p><b>Remediation:</b> {(f.remediation or '').replace('<', '&lt;')}</p>
          <p><b>Status:</b> {f.verification_status}{' (suppressed)' if f.suppressed else ''}</p>
        </div>
        """)
    summary = scan.summary or {}
    html = f"""
    <!doctype html>
    <html><head><style>
      body {{ font-family: Helvetica, Arial, sans-serif; color:#111; }}
      h1 {{ border-bottom: 4px solid #000; padding-bottom: 8px; }}
      .grade {{ display:inline-block; padding: 10px 30px; border:3px solid #000;
               font-size:48px; font-weight:900; background:#FFD23F; box-shadow:6px 6px 0 #000; }}
      table {{ border-collapse:collapse; margin: 16px 0; }}
      th, td {{ border:2px solid #000; padding:6px 12px; }}
      .finding {{ margin:16px 0; padding:12px; border:3px solid #000;
                 box-shadow:4px 4px 0 #000; background:#FDFBF5; }}
    </style></head><body>
      <h1>Pencheff Security Assessment</h1>
      <p><b>Target:</b> {target.name} — {target.base_url}<br>
         <b>Scan ID:</b> {scan.id}<br>
         <b>Generated:</b> {datetime.now(timezone.utc).isoformat()}</p>
      <p class="grade">{scan.grade or '-'}</p>
      <p><b>Score:</b> {scan.score or 0}/100</p>
      <h2>Summary</h2>
      <table>
        <tr><th>Critical</th><th>High</th><th>Medium</th><th>Low</th><th>Info</th></tr>
        <tr>
          <td>{summary.get('critical', 0)}</td>
          <td>{summary.get('high', 0)}</td>
          <td>{summary.get('medium', 0)}</td>
          <td>{summary.get('low', 0)}</td>
          <td>{summary.get('info', 0)}</td>
        </tr>
      </table>
      <h2>Findings</h2>
      {''.join(rows)}
      <hr style="margin-top:32px;border:none;border-top:1px solid #888;">
      <p style="font-size:10px;color:#555;font-style:italic;">
        OWASP and OWASP Top 10 are trademarks of the OWASP Foundation.
        Pencheff is not affiliated with or endorsed by OWASP. Other framework
        names referenced in this report are the property of their respective
        owners and are used for identification only.
      </p>
    </body></html>
    """
    HTML(string=html).write_pdf(str(path))
    return path.stat().st_size


# ── Markdown export (consultancy-friendly) ────────────────────────────

def write_markdown(scan: Scan, target: Target, findings: list[Finding], path: Path,
                   branding: WorkspaceBranding | None = None,
                   threat_model: dict | None = None) -> int:
    lines: list[str] = []
    if branding and branding.logo_url:
        lines.append(f"![logo]({branding.logo_url})\n")
    lines.append(f"# Pencheff Security Assessment — {target.name}\n")
    if branding and branding.opening_letter_md:
        lines.append(branding.opening_letter_md.strip() + "\n")
    lines.append(f"- **Target**: {target.base_url}")
    lines.append(f"- **Scan ID**: `{scan.id}`")
    lines.append(f"- **Profile**: `{scan.profile}`")
    if scan.grade:
        lines.append(f"- **Grade**: **{scan.grade}** ({scan.score or 0}/100)")
    lines.append(f"- **Generated**: {datetime.now(timezone.utc).isoformat()}\n")
    if branding and branding.methodology_md:
        lines.append("## Methodology\n")
        lines.append(branding.methodology_md.strip() + "\n")
    summary = scan.summary or {}
    if threat_model:
        from .threat_model import render_markdown as _render_threat_model
        # Render the engagement-scoped STRIDE/DREAD model inline. The
        # service emits its own ``# STRIDE/DREAD Threat Model`` heading,
        # so we wrap it in a section anchor for the renderer's TOC.
        lines.append("## Threat model")
        lines.append("")
        lines.append(_render_threat_model(threat_model).strip())
        lines.append("")
    lines.append("## Summary\n")
    lines.append("| Critical | High | Medium | Low | Info |")
    lines.append("|---|---|---|---|---|")
    lines.append(
        f"| {summary.get('critical', 0)} | {summary.get('high', 0)} | "
        f"{summary.get('medium', 0)} | {summary.get('low', 0)} | "
        f"{summary.get('info', 0)} |\n"
    )
    lines.append("## Findings\n")
    for f in _sort_findings(findings):
        lines.append(f"### {f.severity.upper()} — {f.title}")
        lines.append(f"- **OWASP**: {f.owasp_category or '-'}  **CWE**: {f.cwe_id or '-'}")
        if f.cvss_score:
            lines.append(f"- **CVSS**: {f.cvss_score}")
        lines.append(f"- **Endpoint**: `{f.endpoint or '-'}`  **Parameter**: `{f.parameter or '-'}`")
        lines.append(f"- **Status**: {f.verification_status}")
        lines.append("")
        if f.description:
            lines.append(f.description.strip() + "\n")
        if f.remediation:
            lines.append(f"**Remediation**: {f.remediation.strip()}\n")
    if branding and branding.footer_text:
        lines.append("---\n")
        lines.append(branding.footer_text)
    payload = "\n".join(lines)
    path.write_text(payload, encoding="utf-8")
    return len(payload.encode())


# ── Delta report ──────────────────────────────────────────────────────

def write_delta_markdown(
    scan_a: Scan, scan_b: Scan, target: Target,
    findings_a: list[Finding], findings_b: list[Finding], path: Path,
    branding: WorkspaceBranding | None = None,
) -> int:
    """Emit a Markdown delta report comparing scan_a (older) to scan_b (newer).

    A finding is keyed by (title, endpoint, parameter, owasp_category) for
    the diff. New = present in B not A. Fixed = present in A not B.
    Regressed = present in both but severity worsened.
    """
    def _key(f: Finding) -> tuple:
        return (f.title, f.endpoint or "", f.parameter or "", f.owasp_category or "")

    a_index = {_key(f): f for f in findings_a}
    b_index = {_key(f): f for f in findings_b}
    new = [f for k, f in b_index.items() if k not in a_index]
    fixed = [f for k, f in a_index.items() if k not in b_index]
    regressed = [
        f for k, f in b_index.items()
        if k in a_index and SEVERITY_ORDER.get(f.severity, 99) < SEVERITY_ORDER.get(a_index[k].severity, 99)
    ]

    lines: list[str] = []
    if branding and branding.logo_url:
        lines.append(f"![logo]({branding.logo_url})\n")
    lines.append(f"# Re-test report — {target.name}\n")
    lines.append(f"Comparing scan `{scan_a.id[:8]}` ({scan_a.created_at.isoformat()}) → "
                 f"`{scan_b.id[:8]}` ({scan_b.created_at.isoformat()}).\n")
    lines.append(f"- **New findings**: {len(new)}")
    lines.append(f"- **Fixed findings**: {len(fixed)}")
    lines.append(f"- **Regressed (severity ↑)**: {len(regressed)}\n")

    def _section(title: str, rows: list[Finding]):
        lines.append(f"## {title}\n")
        if not rows:
            lines.append("_none_\n")
            return
        for f in _sort_findings(rows):
            lines.append(f"- **{f.severity.upper()}** — {f.title} (`{f.endpoint or '-'}`)")
        lines.append("")

    _section("New", new)
    _section("Fixed", fixed)
    _section("Regressed", regressed)
    payload = "\n".join(lines)
    path.write_text(payload, encoding="utf-8")
    return len(payload.encode())


# ── Orchestrator ──────────────────────────────────────────────────────

WRITERS = {"json": write_json, "csv": write_csv, "docx": write_docx, "pdf": write_pdf, "markdown": write_markdown}
EXT = {"json": "json", "csv": "csv", "docx": "docx", "pdf": "pdf", "markdown": "md"}


async def generate_report(report_id: str) -> None:
    settings = get_settings()
    engine = create_async_engine(settings.database_url, future=True)
    Session = async_sessionmaker(engine, expire_on_commit=False, class_=AsyncSession)

    async with Session() as db:
        report = (await db.execute(select(Report).where(Report.id == report_id))).scalar_one_or_none()
        if not report:
            return
        scan = (await db.execute(select(Scan).where(Scan.id == report.scan_id))).scalar_one()
        target = (await db.execute(select(Target).where(Target.id == scan.target_id))).scalar_one()
        findings = (await db.execute(
            select(Finding).where(Finding.scan_id == scan.id).where(Finding.suppressed.is_(False))
        )).scalars().all()
        branding = (await db.execute(
            select(WorkspaceBranding).where(WorkspaceBranding.workspace_id == scan.workspace_id)
        )).scalar_one_or_none()
        # If the scan was scoped to an engagement that has a threat model
        # attached, pull it so write_markdown can include a "Threat model"
        # section. Reports for ad-hoc scans (no engagement) skip this.
        threat_model: dict | None = None
        if scan.engagement_id:
            eng = (await db.execute(
                select(Engagement).where(Engagement.id == scan.engagement_id)
            )).scalar_one_or_none()
            if eng is not None and eng.threat_model:
                threat_model = eng.threat_model
        compared = None
        compared_findings: list[Finding] = []
        if report.kind == "delta" and report.compared_scan_id:
            compared = (await db.execute(select(Scan).where(Scan.id == report.compared_scan_id))).scalar_one_or_none()
            if compared is not None:
                compared_findings = (await db.execute(
                    select(Finding).where(Finding.scan_id == compared.id).where(Finding.suppressed.is_(False))
                )).scalars().all()

    fmt = report.format
    writer = WRITERS.get(fmt)
    if not writer and report.kind != "delta":
        async with Session() as db:
            r = (await db.execute(select(Report).where(Report.id == report_id))).scalar_one()
            r.status = "failed"
            await db.commit()
        return

    out_dir = _storage_dir() / scan.id
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = f"pencheff_{scan.id[:8]}_{_timestamp()}.{EXT.get(fmt, 'md')}"
    path = out_dir / filename

    try:
        if report.kind == "delta" and compared is not None:
            size = write_delta_markdown(
                compared, scan, target, compared_findings, findings, path, branding,
            )
        elif fmt == "markdown":
            size = write_markdown(
                scan, target, findings, path, branding, threat_model=threat_model
            )
        else:
            size = writer(scan, target, findings, path)
    except Exception as e:
        log.exception("report generation failed")
        async with Session() as db:
            r = (await db.execute(select(Report).where(Report.id == report_id))).scalar_one()
            r.status = "failed"
            await db.commit()
        return

    async with Session() as db:
        r = (await db.execute(select(Report).where(Report.id == report_id))).scalar_one()
        r.status = "ready"
        r.storage_path = str(path)
        r.bytes = size
        r.generated_at = datetime.now(timezone.utc)
        await db.commit()


def generate_report_sync(report_id: str) -> None:
    asyncio.run(generate_report(report_id))
