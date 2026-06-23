"""Export pentest results to Word (.docx), CSV, and JSON files."""

from __future__ import annotations

import csv
import io
import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from pencheff.config import Severity, VerificationStatus
from pencheff.core.session import PentestSession
from pencheff.reporting.compliance import get_compliance_summary, get_owasp_coverage
from pencheff.reporting.renderer import render_report
from pencheff.modules.llm_red_team.reporting import (
    build_red_team_summary,
    render_junit_xml,
    render_prometheus_metrics,
)


def _output_dir(session: PentestSession, output_dir: str | None) -> Path:
    """Resolve output directory — default to ~/pencheff-reports/<session_id>/."""
    if output_dir:
        p = Path(output_dir)
    else:
        p = Path.home() / "pencheff-reports" / session.id
    p.mkdir(parents=True, exist_ok=True)
    return p


def _timestamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


# ── JSON Export ──────────────────────────────────────────────────────


def export_json(session: PentestSession, output_dir: str | None = None) -> str:
    """Export findings to a detailed JSON file with verification status."""
    out = _output_dir(session, output_dir)
    findings = session.findings.get_all()
    summary = session.findings.summary()

    data = {
        "report_metadata": {
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "target": session.target.base_url,
            "session_id": session.id,
            "depth": session.depth.value,
            "total_findings": session.findings.count,
            "total_requests": len(session.request_log),
        },
        "summary": {
            "by_severity": summary,
            "by_verification_status": _verification_summary(findings),
        },
        "attack_surface": {
            "endpoints_discovered": len(session.discovered.endpoints),
            "subdomains_found": len(session.discovered.subdomains),
            "open_ports": len(session.discovered.open_ports),
            "tech_stack": session.discovered.tech_stack,
        },
        "findings": [],
        "compliance": get_compliance_summary(findings, [
            "owasp", "pci-dss", "nist", "soc2", "iso27001", "hipaa"
        ]),
        "suppressed_findings": [
            f.to_dict() for f in session.findings.get_all(include_suppressed=True)
            if f.suppressed
        ],
    }
    redteam_summary = build_red_team_summary(findings)
    if redteam_summary["total_failures"]:
        data["llm_redteam"] = redteam_summary

    for f in findings:
        entry = f.to_dict()
        # Flatten compliance for readability
        entry["pci_dss"] = entry.get("compliance", {}).get("PCI-DSS", [])
        entry["nist_800_53"] = entry.get("compliance", {}).get("NIST-800-53", [])
        entry["owasp_mapping"] = entry.get("compliance", {}).get("OWASP", [])
        data["findings"].append(entry)

    filepath = out / f"pencheff_findings_{_timestamp()}.json"
    filepath.write_text(json.dumps(data, indent=2, default=str), encoding="utf-8")
    return str(filepath)


# ── CSV Export ───────────────────────────────────────────────────────


def export_csv(session: PentestSession, output_dir: str | None = None) -> str:
    """Export findings to CSV with verification status and compliance details."""
    out = _output_dir(session, output_dir)
    findings = session.findings.get_all()

    filepath = out / f"pencheff_findings_{_timestamp()}.csv"

    fieldnames = [
        "id",
        "title",
        "severity",
        "cvss_score",
        "cvss_vector",
        "category",
        "owasp_category",
        "owasp_name",
        "endpoint",
        "parameter",
        "description",
        "remediation",
        "cwe_id",
        "verification_status",
        "verification_notes",
        "suppressed",
        "suppress_reason",
        "suppress_notes",
        "evidence_count",
        "evidence_summary",
        "pci_dss",
        "nist_800_53",
        "soc2",
        "iso27001",
        "hipaa",
        "references",
        "discovered_at",
    ]

    with open(filepath, "w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()

        for f in findings:
            evidence_parts = []
            for e in f.evidence[:5]:
                desc = e.description or ""
                evidence_parts.append(f"{e.request_method} {e.request_url} [{e.response_status}] {desc}")

            compliance = f.compliance_mapping
            writer.writerow({
                "id": f.id,
                "title": f.title,
                "severity": f.severity.value,
                "cvss_score": f.cvss_score,
                "cvss_vector": f.cvss_vector,
                "category": f.category,
                "owasp_category": f.owasp_category,
                "owasp_name": f.owasp_name,
                "endpoint": f.endpoint,
                "parameter": f.parameter or "",
                "description": f.description,
                "remediation": f.remediation,
                "cwe_id": f.cwe_id or "",
                "verification_status": f.verification_status.value,
                "verification_notes": f.verification_notes,
                "suppressed": str(f.suppressed),
                "suppress_reason": f.suppress_reason.value if f.suppress_reason else "",
                "suppress_notes": f.suppress_notes,
                "evidence_count": len(f.evidence),
                "evidence_summary": " | ".join(evidence_parts),
                "pci_dss": ", ".join(compliance.get("PCI-DSS", [])),
                "nist_800_53": ", ".join(compliance.get("NIST-800-53", [])),
                "soc2": ", ".join(compliance.get("SOC 2", [])),
                "iso27001": ", ".join(compliance.get("ISO 27001:2022", [])),
                "hipaa": ", ".join(compliance.get("HIPAA", [])),
                "references": ", ".join(f.references),
                "discovered_at": f.discovered_at.isoformat(),
            })

    return str(filepath)


def export_junit(session: PentestSession, output_dir: str | None = None) -> str:
    """Export LLM red-team findings to JUnit XML for CI systems."""
    out = _output_dir(session, output_dir)
    findings = session.findings.get_all()
    filepath = out / f"pencheff_llm_redteam_{_timestamp()}.xml"
    filepath.write_text(render_junit_xml(findings), encoding="utf-8")
    return str(filepath)


def export_prometheus(session: PentestSession, output_dir: str | None = None) -> str:
    """Export LLM red-team metrics in Prometheus text format."""
    out = _output_dir(session, output_dir)
    findings = session.findings.get_all()
    filepath = out / f"pencheff_llm_redteam_{_timestamp()}.prom"
    filepath.write_text(render_prometheus_metrics(findings), encoding="utf-8")
    return str(filepath)


# ── Word (.docx) Export ──────────────────────────────────────────────


def export_docx(
    session: PentestSession,
    report_type: str = "full",
    output_dir: str | None = None,
) -> str:
    """Export penetration test report to Word (.docx) format."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    out = _output_dir(session, output_dir)
    findings = session.findings.get_all()
    summary = session.findings.summary()

    doc = Document()

    # ── Styles ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── Title Page ──
    title = doc.add_heading("Penetration Test Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Light Grid Accent 1"
    meta_rows = [
        ("Target", session.target.base_url),
        ("Date", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
        ("Session ID", session.id),
        ("Test Depth", session.depth.value),
        ("Total Findings", str(session.findings.count)),
        ("Total Requests", str(len(session.request_log))),
    ]
    for i, (label, value) in enumerate(meta_rows):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = value
        for cell in meta_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                paragraph.style = doc.styles["Normal"]

    doc.add_page_break()

    # ── Executive Summary ──
    if report_type in ("executive", "full"):
        doc.add_heading("Executive Summary", level=1)

        # Severity table
        sev_table = doc.add_table(rows=6, cols=2)
        sev_table.style = "Light Grid Accent 1"
        sev_table.rows[0].cells[0].text = "Severity"
        sev_table.rows[0].cells[1].text = "Count"
        _bold_row(sev_table.rows[0])

        sev_colors = {
            Severity.CRITICAL: RGBColor(0xCC, 0x00, 0x00),
            Severity.HIGH: RGBColor(0xFF, 0x66, 0x00),
            Severity.MEDIUM: RGBColor(0xFF, 0xCC, 0x00),
            Severity.LOW: RGBColor(0x00, 0x99, 0x33),
            Severity.INFO: RGBColor(0x33, 0x66, 0xCC),
        }

        for i, sev in enumerate([Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO], 1):
            count = summary.get(sev.value, 0)
            sev_table.rows[i].cells[0].text = sev.value.upper()
            sev_table.rows[i].cells[1].text = str(count)
            # Color the severity label
            for paragraph in sev_table.rows[i].cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = sev_colors[sev]
                    run.bold = True

        doc.add_paragraph("")

        # Verification status summary
        v_summary = _verification_summary(findings)
        doc.add_heading("Verification Summary", level=2)
        v_table = doc.add_table(rows=len(v_summary) + 1, cols=2)
        v_table.style = "Light Grid Accent 1"
        v_table.rows[0].cells[0].text = "Verification Status"
        v_table.rows[0].cells[1].text = "Count"
        _bold_row(v_table.rows[0])
        for i, (status, count) in enumerate(v_summary.items(), 1):
            v_table.rows[i].cells[0].text = status.replace("_", " ").title()
            v_table.rows[i].cells[1].text = str(count)

        doc.add_paragraph("")

        # Risk assessment
        critical = summary.get("critical", 0)
        high = summary.get("high", 0)
        if critical > 0:
            risk = "CRITICAL"
            assessment = f"The application has {critical} critical vulnerability(ies) requiring immediate remediation."
        elif high > 0:
            risk = "HIGH"
            assessment = f"The application has {high} high-severity vulnerability(ies) that should be addressed promptly."
        elif summary.get("medium", 0) > 0:
            risk = "MEDIUM"
            assessment = "The application has moderate security issues that should be addressed in the near term."
        else:
            risk = "LOW"
            assessment = "No critical or high-severity issues were found. Continue regular security assessments."

        risk_para = doc.add_paragraph()
        run = risk_para.add_run(f"Overall Risk Level: {risk}")
        run.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph(assessment)

        # Attack surface
        doc.add_heading("Attack Surface", level=2)
        doc.add_paragraph(f"Endpoints discovered: {len(session.discovered.endpoints)}")
        doc.add_paragraph(f"Subdomains found: {len(session.discovered.subdomains)}")
        doc.add_paragraph(f"Open ports: {len(session.discovered.open_ports)}")

        # Tech stack
        if session.discovered.tech_stack:
            doc.add_heading("Technology Stack", level=2)
            for category, techs in session.discovered.tech_stack.items():
                doc.add_paragraph(f"{category}: {', '.join(techs)}", style="List Bullet")

        doc.add_page_break()

    # ── Detailed Findings ──
    if report_type in ("technical", "full"):
        doc.add_heading("Detailed Findings", level=1)

        for sev in [Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM, Severity.LOW, Severity.INFO]:
            sev_findings = [f for f in findings if f.severity == sev]
            if not sev_findings:
                continue

            doc.add_heading(f"{sev.value.upper()} ({len(sev_findings)})", level=2)

            for f in sev_findings:
                doc.add_heading(f"[{f.id}] {f.title}", level=3)

                # Finding details table
                detail_rows = [
                    ("CVSS Score", f"{f.cvss_score} ({f.cvss_vector})"),
                    ("Category", f.category),
                    ("OWASP", f"{f.owasp_category}: {f.owasp_name}"),
                    ("Endpoint", f.endpoint),
                    ("Verification Status", f.verification_status.value.replace("_", " ").title()),
                ]
                if f.parameter:
                    detail_rows.append(("Parameter", f.parameter))
                if f.cwe_id:
                    detail_rows.append(("CWE", f.cwe_id))
                if f.verification_notes:
                    detail_rows.append(("Verification Notes", f.verification_notes))

                detail_table = doc.add_table(rows=len(detail_rows), cols=2)
                detail_table.style = "Light Grid Accent 1"
                for i, (label, value) in enumerate(detail_rows):
                    detail_table.rows[i].cells[0].text = label
                    detail_table.rows[i].cells[1].text = value

                doc.add_paragraph("")
                doc.add_paragraph(f"Description: {f.description}")

                if f.evidence:
                    doc.add_paragraph("Evidence:", style="Heading 4")
                    for e in f.evidence[:3]:
                        ev_text = f"{e.request_method} {e.request_url}"
                        if e.description:
                            ev_text += f" - {e.description}"
                        if e.response_body_snippet:
                            ev_text += f"\nResponse: {e.response_body_snippet[:200]}"
                        doc.add_paragraph(ev_text, style="List Bullet")

                doc.add_paragraph(f"Remediation: {f.remediation}")

                if f.references:
                    doc.add_paragraph("References:", style="Heading 4")
                    for ref in f.references:
                        doc.add_paragraph(ref, style="List Bullet")

                # Compliance mapping
                compliance = f.compliance_mapping
                if compliance:
                    parts = []
                    for fw, reqs in compliance.items():
                        parts.append(f"{fw}: {', '.join(reqs)}")
                    doc.add_paragraph(f"Compliance: {' | '.join(parts)}")

                doc.add_paragraph("")  # spacer

    # ── Compliance Summary ──
    if report_type in ("executive", "full"):
        doc.add_page_break()
        doc.add_heading("Compliance Summary", level=1)

        categories = list(set(f.category for f in findings))
        coverage = get_owasp_coverage(categories)

        doc.add_heading("OWASP Top 10 Coverage", level=2)
        owasp_table = doc.add_table(rows=len(coverage) + 1, cols=3)
        owasp_table.style = "Light Grid Accent 1"
        owasp_table.rows[0].cells[0].text = "Category"
        owasp_table.rows[0].cells[1].text = "Tested"
        owasp_table.rows[0].cells[2].text = "Findings"
        _bold_row(owasp_table.rows[0])

        for i, (cat, tested) in enumerate(coverage.items(), 1):
            cat_code = cat.split(":")[0]
            count = sum(1 for f in findings if f.owasp_category == cat_code)
            owasp_table.rows[i].cells[0].text = cat
            owasp_table.rows[i].cells[1].text = "Yes" if tested else "No"
            owasp_table.rows[i].cells[2].text = str(count)

    # ── Remediation Roadmap ──
    if report_type in ("executive", "full"):
        doc.add_page_break()
        doc.add_heading("Remediation Roadmap", level=1)

        doc.add_heading("Immediate (Critical/High)", level=2)
        for f in findings:
            if f.severity in (Severity.CRITICAL, Severity.HIGH):
                doc.add_paragraph(f"{f.title} - {f.remediation}", style="List Bullet")

        doc.add_heading("Short-term (Medium)", level=2)
        for f in findings:
            if f.severity == Severity.MEDIUM:
                doc.add_paragraph(f"{f.title} - {f.remediation}", style="List Bullet")

        doc.add_heading("Long-term (Low/Info)", level=2)
        for f in findings:
            if f.severity in (Severity.LOW, Severity.INFO):
                doc.add_paragraph(f"{f.title} - {f.remediation}", style="List Bullet")

    # ── Footer ──
    doc.add_paragraph("")
    footer = doc.add_paragraph("Report generated by Pencheff")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.italic = True
    footer.runs[0].font.size = Pt(9)

    filepath = out / f"pencheff_report_{_timestamp()}.docx"
    doc.save(str(filepath))
    return str(filepath)


# ── Export All ───────────────────────────────────────────────────────


def export_all(
    session: PentestSession,
    report_type: str = "full",
    output_dir: str | None = None,
) -> dict[str, str]:
    """Export report in all formats: Word (.docx), CSV, and JSON."""
    out_str = str(_output_dir(session, output_dir))
    return {
        "docx": export_docx(session, report_type=report_type, output_dir=out_str),
        "csv": export_csv(session, output_dir=out_str),
        "json": export_json(session, output_dir=out_str),
    }


# ── Helpers ──────────────────────────────────────────────────────────


def _verification_summary(findings: list) -> dict[str, int]:
    """Count findings by verification status."""
    counts: dict[str, int] = {}
    for status in VerificationStatus:
        counts[status.value] = 0
    for f in findings:
        counts[f.verification_status.value] += 1
    return counts


def _bold_row(row) -> None:
    """Make all text in a table row bold."""
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
